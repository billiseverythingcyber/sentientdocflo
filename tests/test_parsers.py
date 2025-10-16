from io import BytesIO
from docx import Document
import pdfplumber

def test_docx_basic_read():
# make a tiny DOCX in memory
doc = Document()
doc.add_paragraph("Hello Invoice #: INV-0001")
b = BytesIO()
doc.save(b)
b.seek(0)
# open and read text
doc2 = Document(b)
text = "
".join(p.text for p in doc2.paragraphs)
assert "INV-0001" in text


def test_pdf_library_imports():
# simple sanity check that pdfplumber is available
assert hasattr(pdfplumber, "__version__")